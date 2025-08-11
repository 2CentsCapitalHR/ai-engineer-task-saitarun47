import os
import re
import io
import json
import tempfile
from pathlib import Path
from docx import Document
from typing import List, Dict, Any
from docx.enum.text  import WD_COLOR_INDEX
from docx.shared import RGBColor

from agno.agent import Agent
from agno.models.google import Gemini
from agno.knowledge.pdf import PDFKnowledgeBase
from agno.knowledge.pdf import PDFImageReader
from agno.tools.knowledge import KnowledgeTools
from agno.vectordb.lancedb import LanceDb
from agno.document.chunking.agentic import AgenticChunking
from sentence_transformers import SentenceTransformer
from agno.embedder.sentence_transformer import SentenceTransformerEmbedder

from dotenv import load_dotenv
load_dotenv()




db_uri = Path(__file__).parent / "lancedb"
lance_table="adgm_regs"






def llm_json(text:str):
    try:
        return json.loads(text)
    except Exception:
        m = re.search(r"(\{.*\}|\[.*\])", text, flags=re.S)
        if m:
            try:
                return json.loads(m.group(1))
            except Exception:
                return {"raw": text}
        
    return {"raw": text}

def build_analysis_prompt(doc_text: str) -> str:
    return f"""
    You are an ADGM Corporate Agent reviewing a user-submitted document for compliance with ADGM Companies Regulations.
    Use ONLY knowledge retrieved from the ADGM regulations. If unsure, say so.

    Return ONLY a JSON array of issues. Each issue must have:
    - section: short locator (or 'General')
    - issue: concise description
    - severity: High | Medium | Low
    - quote: exact problematic snippet if possible
    - suggestion: ADGM-aligned fix
    - rationale: brief 'why' under ADGM rules
    - citation: <=240 chars excerpt from retrieved text (no invented article numbers)
    - source_title: document title of the retrieved source
    - locator_hint: heading from the retrieved source (e.g., "p.2 — Evidence of Appointment")

    Focus on:
    - Incorrect jurisdiction (e.g., UAE Federal Courts vs ADGM Courts/arbitration)
    - Ambiguous language where 'shall' is required
    - Missing signatory/execution blocks and dates
    - Registered office, company name suffixes, directors requirements (where relevant)
    - Incorporation-related clauses

    Document:
    \"\"\"{doc_text}\"\"\"
    """


def analyze_with_agno(agent: Agent, doc_text: str) -> List[Dict[str, Any]]:
    prompt = build_analysis_prompt(doc_text)
    res = agent.run(prompt)
    text = getattr(res, "content", None) or getattr(res, "text", None) or str(res)
    parsed = llm_json(text)
    if isinstance(parsed,list):
        return parsed
    if isinstance(parsed,dict) and "issues" in parsed:
        return parsed["issues"]
    return[]



def add_inline_comments(doc_bytes : bytes, issues :List[Dict[str, Any]]) -> bytes:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        tf.write(doc_bytes)
        path = tf.name
    doc = Document(path)

    def annotate_par(par, note_text):
        for run in par.runs:
            run.font.highlight_color=WD_COLOR_INDEX.YELLOW
        r = par.add_run(f"[COMMENT: {note_text}]")
        r.font.bold=True
        r.font.color.rgb=RGBColor(0xB2, 0x22, 0x22)

    for issue in issues or []:
        quote = (issue.get("quote") or "").strip()
        note = f"{issue.get('severity','')}: {issue.get('issue','')}. Suggestion: {issue.get('suggestion','')}. Ref: {issue.get('citation','')}"
        placed = False

        if quote and len(quote) > 8:
            for par in doc.paragraphs:
                if quote in par.text:
                    annotate_par(par,note)
                    placed=True
                    break

        if not placed:
            lower_issue = (issue.get("issue","") or "").lower()
            keys = []
            if "jurisdiction" in lower_issue:
                keys = ["jurisdiction", "governing law", "court", "arbitration"]
            elif "signat" in lower_issue or "execut" in lower_issue:
                keys = ["sign", "signature", "execution", "date"]
            elif "registered office" in lower_issue or "address" in lower_issue:
                keys = ["registered office", "address"]
            for par in doc.paragraphs:
                if any(k in par.text.lower() for k in keys):
                    annotate_par(par, note)
                    placed = True
                    break

        if not placed and doc.paragraphs:
            annotate_par(doc.paragraphs[-1], note)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    try:
        os.remove(path)
    except Exception:
        pass
    return out.getvalue()


### Agent

def build_agent() -> Agent:
    embedder = SentenceTransformerEmbedder(id="sentence-transformers/all-MiniLM-L6-v2")
    embedder.sentence_transformer_client = SentenceTransformer(embedder.id,device="cpu")

    vector_db = LanceDb(
        table_name=lance_table,
        uri=db_uri,
        embedder=embedder,
    )


    knowledge_base = PDFKnowledgeBase(
        path="data/pdfs",
        reader=PDFImageReader(),
        vector_db=vector_db,
        chunking_strategy=AgenticChunking(model=Gemini(id="gemini-2.0-flash")),
    )

    #knowledge_base.load(recreate=False)  ## comment this after 1st run

    knowledge_tools = KnowledgeTools(
        knowledge=knowledge_base,
        think=True,
        search=True,
        analyze=True,
        add_few_shot=True,

    )

    agent = Agent(
        model=Gemini(id="gemini-2.0-flash"),
        knowledge=knowledge_base,
        search_knowledge=True,
        tools=[knowledge_tools],
        show_tool_calls=True,
        markdown=True,
        debug_mode=True,
    )

    return agent
