import os
import re
import tempfile
from typing import List 
from docx import Document


incorporation_checklist=[
    "Articles of Association",
    "Memorandum of Association",
    "Incorporation Application Form",
    "UBO Declaration Form",
    "Register of Members and Directors",
]

doc_type_patterns=[
    (r"\barticles of association\b|\baoa\b", "Articles of Association"),
    (r"\bmemorandum of association\b|\bmoa\b|\bmou\b", "Memorandum of Association"),
    (r"\bincorporation application\b|\bapplication form\b", "Incorporation Application Form"),
    (r"\bubo\b|\bultimate beneficial owner\b", "UBO Declaration Form"),
    (r"\bregister of members\b|\bregister of directors\b", "Register of Members and Directors"),
    (r"\bboard resolution\b", "Board Resolution"),
    (r"\bshareholder resolution\b", "Shareholder Resolution"),
    (r"\bchange of registered address\b|\bchange of address\b", "Change of Registered Address Notice"),
]






def extract_text(file_bytes: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        tf.write(file_bytes)
        path = tf.name
    doc = Document(path)
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for t in doc.tables:
        for row in t.rows :
            for cell in row.cells:
                texts.append(cell.text)

    try:
        os.remove(path)
    except Exception:
        pass

    return "\n".join(texts)



def classify_doc_type(filename:str , content:str)-> str:
    hay = f"{filename}\n{content}".lower()
    for pattern , label in doc_type_patterns:
        if re.search(pattern,hay):
            return label
    return

def infer_process(doc_types: List[str]) -> str:
    must = {"Articles of Association", "Memorandum of Association"}
    if any(dt in doc_types for dt in must) or "Incorporation Application Form" in doc_types:
        return "Company Incorporation"
    return "Unknown"

def checklist_for_process(process: str) -> List[str]:
    if process == "Company Incorporation":
        return incorporation_checklist
    return []

